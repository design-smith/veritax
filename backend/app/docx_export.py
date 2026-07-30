"""Build a native .docx (cover, headings, prose, native tables, native EDITABLE charts) in pure Python —
no binary, no AGPL. Charts are real Word charts: a hand-authored DrawingML chart part + an openpyxl
embedded workbook, wired into python-docx via OPC relationships. Bar/column/line/pie.
"""

from __future__ import annotations

import io
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CT_CHART = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
CT_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
RT_CHART = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
RT_PACKAGE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/package"

_MARKER = re.compile(r"\[\[(table|chart):([^\]]+)\]\]")
_COL = [chr(c) for c in range(ord("B"), ord("Z") + 1)]  # data columns start at B (A = categories)


# ── Cover + footer ───────────────────────────────────────────────────────────
def _center(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _run(run, *, size: int, color: str = "111111", bold: bool = False, caps: bool = False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if caps:
        run._element.get_or_add_rPr().append(OxmlElement("w:caps"))
    return run


def _cover(doc, cover: dict) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.5)

    brand = _center(doc.add_paragraph())
    brand.paragraph_format.space_after = Pt(72)
    _run(brand.add_run("VERITAX"), size=10, color="777777", caps=True)

    title = _center(doc.add_paragraph())
    title.paragraph_format.space_after = Pt(18)
    _run(title.add_run(cover.get("documentTitle", "Transfer Pricing Local File")), size=30, color="111111")

    status = _center(doc.add_paragraph())
    status.paragraph_format.space_after = Pt(72)
    _run(status.add_run(cover.get("status", "Draft prepared for review")), size=10, color="777777")

    entity = _center(doc.add_paragraph())
    entity.paragraph_format.space_after = Pt(8)
    _run(entity.add_run("Entity"), size=9, color="777777", caps=True)
    entity.add_run("\n")
    _run(entity.add_run(cover.get("entity") or "Entity"), size=16, color="111111")

    jurisdiction = _center(doc.add_paragraph())
    jurisdiction.paragraph_format.space_after = Pt(72)
    _run(jurisdiction.add_run("Jurisdiction"), size=9, color="777777", caps=True)
    jurisdiction.add_run("\n")
    _run(jurisdiction.add_run(cover.get("jurisdiction", "")), size=14, color="333333")

    footer = _center(doc.add_paragraph())
    prepared = cover.get("preparedBy", "Veritax")
    prepared_on = cover.get("preparedOn")
    line = f"Prepared by {prepared}" + (f" | {prepared_on}" if prepared_on else "")
    _run(footer.add_run(line), size=9, color="888888")
    doc.add_page_break()


def _page_number_footer(doc) -> None:
    footer = doc.sections[0].footer
    p = _center(footer.paragraphs[0])
    p.add_run("Page ")
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(begin); run._r.append(instr); run._r.append(end)


# ── Tables ───────────────────────────────────────────────────────────────────
def _shade(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_table(doc, spec: dict) -> None:
    cols = spec.get("columns", []) or [""]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"  # built-in bordered style → real gridlines
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = str(c)
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
        _shade(cell, "D9D9D9")
    for row in spec.get("rows", []):
        cells = table.add_row().cells
        for i, v in enumerate(row):
            if i < len(cells):
                cells[i].text = str(v)
    doc.add_paragraph()


# ── Native charts (hand-authored DrawingML + embedded workbook) ───────────────
def _cache(ref: str, values: list, numeric: bool) -> str:
    tag = "c:numCache" if numeric else "c:strCache"
    pts = "".join(f'<c:pt idx="{i}"><c:v>{escape(str(v))}</c:v></c:pt>' for i, v in enumerate(values))
    kind = "c:numRef" if numeric else "c:strRef"
    return f"<{kind}><c:f>{ref}</c:f><{tag}><c:ptCount val=\"{len(values)}\"/>{pts}</{tag}></{kind}>"


def _series_xml(idx: int, name: str, name_ref: str, cats: list, cat_ref: str, vals: list, val_ref: str) -> str:
    return (
        f'<c:ser><c:idx val="{idx}"/><c:order val="{idx}"/>'
        f'<c:tx>{_cache(name_ref, [name], False)}</c:tx>'
        f'<c:cat>{_cache(cat_ref, cats, False)}</c:cat>'
        f'<c:val>{_cache(val_ref, vals, True)}</c:val>'
        f'</c:ser>'
    )


def _chart_xml(spec: dict) -> bytes:
    cats = [str(c) for c in spec.get("categories", [])]
    series = spec.get("series", []) or []
    n = len(cats)
    cat_ref = f"Sheet1!$A$2:$A${n + 1}"
    sers = []
    for k, s in enumerate(series):
        col = _COL[k] if k < len(_COL) else "B"
        sers.append(_series_xml(
            k, s.get("name", f"Series {k + 1}"), f"Sheet1!${col}$1",
            cats, cat_ref, s.get("values", []), f"Sheet1!${col}$2:${col}${n + 1}",
        ))
    ser_xml = "".join(sers)
    ctype = spec.get("type", "bar")
    if ctype == "pie":
        plot = f"<c:pieChart><c:varyColors val=\"1\"/>{ser_xml}</c:pieChart>"
    elif ctype == "line":
        plot = (f'<c:lineChart><c:grouping val="standard"/>{ser_xml}'
                f'<c:marker val="1"/><c:axId val="1"/><c:axId val="2"/></c:lineChart>'
                f'{_axes()}')
    else:  # bar / column → vertical columns
        plot = (f'<c:barChart><c:barDir val="col"/><c:grouping val="clustered"/>{ser_xml}'
                f'<c:axId val="1"/><c:axId val="2"/></c:barChart>{_axes()}')
    title = escape(str(spec.get("title", "")))
    title_xml = (f'<c:title><c:tx><c:rich><a:bodyPr/><a:p><a:r><a:t>{title}</a:t></a:r></a:p></c:rich></c:tx>'
                 f'<c:overlay val="0"/></c:title><c:autoTitleDeleted val="0"/>') if title else '<c:autoTitleDeleted val="1"/>'
    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<c:chart>{title_xml}<c:plotArea><c:layout/>{plot}</c:plotArea>'
        '<c:plotVisOnly val="1"/><c:dispBlanksAs val="gap"/></c:chart>'
        '<c:externalData r:id="rId1"><c:autoUpdate val="0"/></c:externalData>'
        '</c:chartSpace>'
    )
    return xml.encode("utf-8")


def _axes() -> str:
    return (
        '<c:catAx><c:axId val="1"/><c:scaling><c:orientation val="minMax"/></c:scaling>'
        '<c:delete val="0"/><c:axPos val="b"/><c:crossAx val="2"/></c:catAx>'
        '<c:valAx><c:axId val="2"/><c:scaling><c:orientation val="minMax"/></c:scaling>'
        '<c:delete val="0"/><c:axPos val="l"/><c:crossAx val="1"/></c:valAx>'
    )


def _chart_xlsx(spec: dict) -> bytes:
    from openpyxl import Workbook
    wb = Workbook(); ws = wb.active; ws.title = "Sheet1"
    cats = spec.get("categories", [])
    series = spec.get("series", []) or []
    for i, s in enumerate(series):
        ws.cell(row=1, column=2 + i, value=s.get("name", f"Series {i + 1}"))
    for r, cat in enumerate(cats, start=2):
        ws.cell(row=r, column=1, value=cat)
        for i, s in enumerate(series):
            vals = s.get("values", [])
            ws.cell(row=r, column=2 + i, value=vals[r - 2] if r - 2 < len(vals) else None)
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()


_EMU_W, _EMU_H = 5486400, 3200400  # 6in × 3.5in


def _add_chart(doc, spec: dict, index: int) -> None:
    package = doc.part.package
    chart_part = Part(PackURI(f"/word/charts/chart{index}.xml"), CT_CHART, _chart_xml(spec), package)
    xlsx_part = Part(PackURI(f"/word/embeddings/Microsoft_Excel_Sheet{index}.xlsx"), CT_XLSX, _chart_xlsx(spec), package)
    chart_part.relate_to(xlsx_part, RT_PACKAGE)          # becomes rId1 (matches externalData)
    rid = doc.part.relate_to(chart_part, RT_CHART)       # document → chart
    drawing = (
        '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<wp:inline distT="0" distB="0" distL="0" distR="0" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        f'<wp:extent cx="{_EMU_W}" cy="{_EMU_H}"/><wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{index}" name="Chart {index}"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">'
        '<c:chart xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{rid}"/>'
        '</a:graphicData></a:graphic></wp:inline></w:drawing>'
    )
    run = doc.add_paragraph().add_run()
    run._r.append(parse_xml(drawing))


# ── Prose + assembly ─────────────────────────────────────────────────────────
def _prose(doc, text: str) -> None:
    for para in re.split(r"\n{2,}", text.strip()):
        para = para.strip()
        if not para:
            continue
        # ponytail: markdown-lite — strip #/*/- lead tokens; full inline styling is a later pass.
        clean = re.sub(r"^\s*#{1,6}\s+", "", para)
        clean = re.sub(r"[*_]{1,2}([^*_]+)[*_]{1,2}", r"\1", clean)
        doc.add_paragraph(clean.replace("\n", " "))


def _render_section(doc, sec: dict, counter: dict) -> None:
    doc.add_heading(sec.get("heading", ""), level=1)
    tables = {t["id"]: t for t in sec.get("tables", []) if "id" in t}
    charts = {c["id"]: c for c in sec.get("charts", []) if "id" in c}
    text = sec.get("content", "") or ""
    pos = 0
    used = set()
    for m in _MARKER.finditer(text):
        pre = text[pos:m.start()]
        if pre.strip():
            _prose(doc, pre)
        kind, cid = m.group(1), m.group(2)
        if kind == "table" and cid in tables:
            _add_table(doc, tables[cid]); used.add(cid)
        elif kind == "chart" and cid in charts:
            counter["c"] += 1; _add_chart(doc, charts[cid], counter["c"]); used.add(cid)
        pos = m.end()
    if text[pos:].strip():
        _prose(doc, text[pos:])
    for t in sec.get("tables", []):           # any not referenced by a marker → append
        if t.get("id") not in used:
            _add_table(doc, t)
    for c in sec.get("charts", []):
        if c.get("id") not in used:
            counter["c"] += 1; _add_chart(doc, c, counter["c"])


def build_document(cover: dict, sections: list[dict]) -> bytes:
    doc = Document()
    _cover(doc, cover)
    _page_number_footer(doc)
    counter = {"c": 0}
    for sec in sections:
        _render_section(doc, sec, counter)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
